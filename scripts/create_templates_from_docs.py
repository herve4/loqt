import os
import django
import sys
import re
import docx

# Configuration de l'environnement Django
sys.path.append('/home/herve/loqt')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'loqt.settings')
django.setup()

from logistque.models import ChronogrammeTemplate
from django.contrib.auth import get_user_model

User = get_user_model()
admin_user = User.objects.filter(is_superuser=True).first()

DOCS = [
    "CHRONOGRAMME ALLOCO PARTY.docx",
    "CHRONOGRAMME DE LA FÊTE DE NOËL DES ENFANTS.docx",
    "CHRONOGRAMME DU CAMP RÉGIONAL DU 07 AU 10 AVRIL 2026 À GRAND BASSAM.docx",
    "CHRONOGRAMME PROGRAMME EN L’HONNEUR DES MAMANS.docx",
    "Chronogramme du 31 déc 2025.docx",
    "PROGRAMME DE DIRECTION-PRIERE DU MOIS DE NOVEMBRE 2025.docx"
]

def parse_line(line):
    """Extrait heure et titre d'une ligne de texte"""
    time_pattern = r'(\d{1,2}[hH:]\d{0,2})(?:\s*[-~à]\s*(\d{1,2}[hH:]\d{0,2}))?'
    match = re.search(time_pattern, line)
    if match:
        start = match.group(1).replace('h', ':').replace('H', ':')
        if ':' not in start: start += ":00"
        end = match.group(2).replace('h', ':').replace('H', ':') if match.group(2) else ""
        if end and ':' not in end: end += ":00"
        
        titre = line.replace(match.group(0), "").strip()
        return {'Heure Début': start, 'Heure Fin': end, 'Activité': titre, 'Détails': '', 'jour': 'Jour 1'}
    return None

def extract_from_docx(path):
    doc = docx.Document(path)
    items = []
    headers = ['Heure Début', 'Heure Fin', 'Activité', 'Détails']
    
    # 1. Tentative Tableaux
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) > 1:
            header_row = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Mode Grille (Lundi, Mardi...)
            days_keywords = ['janv', 'févr', 'mar', 'avr', 'mai', 'juin', 'juil', 'août', 'sept', 'oct', 'nov', 'déc', 'jour']
            is_date_grid = any(any(kw in h.lower() for kw in days_keywords) for h in header_row[1:])
            
            if is_date_grid:
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if not cells[0]: continue
                    time_val = cells[0]
                    for i in range(1, len(cells)):
                        if cells[i] and len(cells[i]) > 2:
                            item = parse_line(f"{time_val} {cells[i]}")
                            if item:
                                item['jour'] = header_row[i]
                                items.append(item)
                if items: return items, headers

            # Mode Vertical standard
            if any(kw in header_row[0].lower() for kw in ['horaire', 'heure', 'time']):
                current_headers = header_row
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if not cells[0]: continue
                    item = {'jour': 'Jour 1'}
                    for i, val in enumerate(cells):
                        label = current_headers[i] if i < len(current_headers) else f"Col {i+1}"
                        item[label] = val
                    items.append(item)
                if items: return items, current_headers

    # 2. Fallback Texte
    if not items:
        for para in doc.paragraphs:
            it = parse_line(para.text.strip())
            if it: items.append(it)
            
    return items, headers

def run():
    for doc_name in DOCS:
        path = os.path.join('/home/herve/loqt', doc_name)
        if not os.path.exists(path):
            print(f"Fichier manquant: {doc_name}")
            continue
            
        print(f"Traitement de {doc_name}...")
        items, headers = extract_from_docx(path)
        
        if items:
            template, created = ChronogrammeTemplate.objects.update_or_create(
                nom=doc_name.replace('.docx', '').title(),
                defaults={
                    'description': f"Modèle généré automatiquement à partir de {doc_name}",
                    'items': items,
                    'headers': headers,
                    'cree_par': admin_user
                }
            )
            print(f"  -> {'Créé' if created else 'Mis à jour'}: {template.nom} ({len(items)} items)")
        else:
            print(f"  -> Aucun item trouvé dans {doc_name}")

if __name__ == "__main__":
    run()
