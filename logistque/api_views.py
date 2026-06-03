from rest_framework import viewsets, permissions, filters, pagination
try:
    import pytesseract
    from PIL import Image
    # Configuration du chemin Tesseract pour le serveur de production
    pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
import re
import datetime
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None
try:
    import pandas as pd
except ImportError:
    pd = None
from rest_framework.decorators import api_view, action
from rest_framework.authentication import SessionAuthentication
from rest_framework.response import Response
from django.contrib.auth import get_user_model
from django.db.models import Count
from django.utils import timezone
import logging

logger = logging.getLogger('logistque')
from django_filters.rest_framework import DjangoFilterBackend
from logistque.models import (
    Region, Ville, Eglise, Materiel, Evenement, ChronogrammeItem, PoleCompetence,
    MouvementMateriel, FicheDefectuosite, ReunionDimanche,
    RessourceFormation, DemandeFormationSGL, ExpressionBesoin,
    ValidationCircuit,
    Formation, DemandeFormation, SessionFormation,
    CategorieMateriel, SousCategorieMateriel, EvenementImage,
    ChronogrammeTemplate, MaterielImage
)

class CsrfExemptSessionAuthentication(SessionAuthentication):
    def enforce_csrf(self, request):
        return  # Ne pas forcer CSRF pour les besoins de l'API
from logistque.serializers import (
    RegionSerializer, VilleSerializer, EgliseSerializer, MaterielSerializer, EvenementSerializer,
    ChronogrammeItemSerializer, PoleCompetenceSerializer,
    MouvementMaterielSerializer, FicheDefectuositeSerializer,
    ReunionDimancheSerializer, RessourceFormationSerializer,
    DemandeFormationSGLSerializer, ExpressionBesoinSerializer,
    ValidationCircuitSerializer,
    FormationSerializer, DemandeFormationSerializer, SessionFormationSerializer,
    CategorieMaterielSerializer, SousCategorieMaterielSerializer,
    EvenementImageSerializer, ChronogrammeTemplateSerializer, MaterielImageSerializer
)

class StandardPagination(pagination.PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 10000

class EgliseViewSet(viewsets.ModelViewSet):
    queryset = Eglise.objects.all().order_by('nom')
    serializer_class = EgliseSerializer
    permission_classes = [permissions.AllowAny]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['region', 'pays']
    search_fields = ['nom', 'phone', 'ville__nom', 'pasteur__email', 'pasteur__first_name', 'pasteur__last_name']
    ordering_fields = ['nom', 'created_at']

class CategorieViewSet(viewsets.ModelViewSet):
    queryset = CategorieMateriel.objects.all().order_by('nom')
    serializer_class = CategorieMaterielSerializer
    permission_classes = [permissions.IsAuthenticated]

class SousCategorieViewSet(viewsets.ModelViewSet):
    queryset = SousCategorieMateriel.objects.all().order_by('nom')
    serializer_class = SousCategorieMaterielSerializer
    permission_classes = [permissions.IsAuthenticated]

class MaterielViewSet(viewsets.ModelViewSet):
    queryset = Materiel.objects.filter(is_deleted=False)
    serializer_class = MaterielSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['categorie', 'sous_categorie', 'eglise', 'etat']
    search_fields = ['nom', 'description', 'identifiant_unique']
    ordering_fields = ['created_at', 'quantite']

    def perform_create(self, serializer):
        materiel = serializer.save()
        images = self.request.FILES.getlist('uploaded_images')
        for img in images:
            MaterielImage.objects.create(materiel=materiel, image=img)
        
        remaining = materiel.images_materiel.all().order_by('id')
        if remaining.exists():
            materiel.image = remaining[0].image
        else:
            materiel.image = None
        materiel.save()

    def perform_update(self, serializer):
        materiel = serializer.save()
        images = self.request.FILES.getlist('uploaded_images')
        for img in images:
            MaterielImage.objects.create(materiel=materiel, image=img)
        
        remaining = materiel.images_materiel.all().order_by('id')
        if remaining.exists():
            materiel.image = remaining[0].image
        else:
            materiel.image = None
        materiel.save()

    def perform_destroy(self, instance):
        # Soft delete
        instance.is_deleted = True
        instance.save()

class EvenementViewSet(viewsets.ModelViewSet):
    queryset = Evenement.objects.all().order_by('-date_debut')
    serializer_class = EvenementSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = pagination.PageNumberPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['type_evenement', 'statut', 'eglise']
    search_fields = ['titre', 'lieu', 'description']
    ordering_fields = ['date_debut', 'created_at']

    @action(detail=False, methods=['post'], url_path='extract-chronogram', authentication_classes=[], permission_classes=[permissions.AllowAny])
    def extract_chronogram(self, request):
        """Extraction intelligente à partir de fichiers PDF, DOCX, XLSX ou Image."""
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'Aucun fichier fourni'}, status=400)
            
        filename = file_obj.name.lower()
        items = []
        raw_text = ""

        try:
            if filename.endswith(('.png', '.jpg', '.jpeg', '.webp')):
                # Traitement Image (OCR)
                img = Image.open(file_obj)
                raw_text = pytesseract.image_to_string(img, lang='eng')
            
            elif filename.endswith('.pdf'):
                # Traitement PDF
                reader = PyPDF2.PdfReader(file_obj)
                for page in reader.pages:
                    raw_text += page.extract_text() + "\n"
            
            elif filename.endswith('.docx'):
                # Traitement Word
                doc = docx.Document(file_obj)
                # Tentative d'extraction structurée si tableaux présents
                for table in doc.tables:
                    if len(table.rows) > 1 and len(table.columns) > 1:
                        header_row = [cell.text.strip() for cell in table.rows[0].cells]
                        
                        # Mode Grille Dynamique
                        if any(kw in header_row[0].lower() for kw in ['horaire', 'heure', 'time']):
                            days_keywords = ['janv', 'févr', 'mar', 'avr', 'mai', 'juin', 'juil', 'août', 'sept', 'oct', 'nov', 'déc', 
                                             'lund', 'mard', 'merc', 'jeud', 'vend', 'sam', 'dim', 'jour']
                            
                            is_date_grid = any(any(kw in h.lower() for kw in days_keywords) for h in header_row[1:])
                            
                            if is_date_grid:
                                # Mode PIVOT : Les colonnes sont des jours
                                days = header_row
                                for row in table.rows[1:]:
                                    cells = [c.text.strip() for c in row.cells]
                                    if not cells[0]: continue
                                    time_val = cells[0]
                                    for i in range(1, len(cells)):
                                        val = cells[i]
                                        if val and len(val) > 2:
                                            day_name = days[i] if i < len(days) else f"Jour {i}"
                                            # Utiliser le parseur de ligne pour extraire heure_debut/fin proprement
                                            parsed = self._parse_line_for_chronogram(f"{time_val} {val}", force_day=day_name)
                                            # Mappage pour le frontend dynamique
                                            for p in parsed:
                                                mapped = {'jour': p.get('jour', 'Jour 1')}
                                                mapped['Heure Début'] = p.get('heure_debut', '')
                                                mapped['Heure Fin'] = p.get('heure_fin', '')
                                                mapped['Activité'] = p.get('titre', '')
                                                mapped['Détails'] = p.get('description', '')
                                                items.append(mapped)
                                if items: return Response({'items': items[:200], 'headers': ['Heure Début', 'Heure Fin', 'Activité', 'Détails'], 'method': 'docx_pivot_grid'})

                            # Mode Grille Standard (une seule journée, colonnes = attributs)
                            headers = header_row
                            for row in table.rows[1:]:
                                cells = [c.text.strip() for c in row.cells]
                                if not cells[0]: continue
                                item = {'jour': 'Jour 1'}
                                for i, val in enumerate(cells):
                                    label = headers[i] if i < len(headers) else f"Col {i+1}"
                                    item[label] = val
                                items.append(item)
                            if items: return Response({'items': items[:100], 'headers': headers, 'method': 'docx_dynamic_grid'})
                        
                        # Mode Vertical Classique avec colonnes dynamiques
                        headers = header_row
                        for row in table.rows[1:]:
                            cells = [c.text.strip() for c in row.cells]
                            if len(cells) >= 2 and re.search(r'\d{1,2}[hH:]', cells[0]):
                                item = {'jour': 'Jour 1'}
                                for i, val in enumerate(cells):
                                    label = headers[i] if i < len(headers) else f"Col {i+1}"
                                    item[label] = val
                                items.append(item)
                        if items: return Response({'items': items[:100], 'headers': headers, 'method': 'docx_dynamic_vertical'})

                # Fallback : texte brut
                for para in doc.paragraphs:
                    raw_text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        raw_text += " | ".join([cell.text for cell in row.cells]) + "\n"
            
            elif filename.endswith(('.xlsx', '.xls')):
                # Traitement Excel
                df = pd.read_excel(file_obj)
                headers = [str(c) for c in df.columns]
                
                days_keywords = ['janv', 'févr', 'mar', 'avr', 'mai', 'juin', 'juil', 'août', 'sept', 'oct', 'nov', 'déc', 
                                 'lund', 'mard', 'merc', 'jeud', 'vend', 'sam', 'dim', 'jour']
                
                is_date_grid = any(any(kw in h.lower() for kw in days_keywords) for h in headers[1:])
                
                if is_date_grid and any(kw in str(headers[0]).lower() for kw in ['horaire', 'heure', 'time']):
                    # Mode PIVOT Excel
                    for _, row in df.iterrows():
                        time_val = str(row.iloc[0])
                        if pd.isna(row.iloc[0]): continue
                        for i in range(1, len(row)):
                            val = str(row.iloc[i])
                            if val and val != 'nan' and len(val) > 2:
                                day_name = headers[i]
                                parsed = self._parse_line_for_chronogram(f"{time_val} {val}", force_day=day_name)
                                for p in parsed:
                                    mapped = {'jour': p.get('jour', 'Jour 1')}
                                    mapped['Heure Début'] = p.get('heure_debut', '')
                                    mapped['Heure Fin'] = p.get('heure_fin', '')
                                    mapped['Activité'] = p.get('titre', '')
                                    mapped['Détails'] = p.get('description', '')
                                    items.append(mapped)
                    if items: return Response({'items': items[:200], 'headers': ['Heure Début', 'Heure Fin', 'Activité', 'Détails'], 'method': 'excel_pivot_grid'})

                # On extrait les données ligne par ligne standard si pas de pivot
                for _, row in df.iterrows():
                    item = {'jour': 'Jour 1'}
                    for i, val in enumerate(row.values):
                        label = headers[i] if i < len(headers) else f"Col {i+1}"
                        item[label] = str(val) if pd.notna(val) else ""
                    items.append(item)
                
                if items: return Response({'items': items[:100], 'headers': headers, 'method': 'excel_dynamic'})
                
                raw_text = df.to_string()
                for _, row in df.iterrows():
                    row_str = " ".join([str(val) for val in row.values if pd.notna(val)])
                    items.extend(self._parse_line_for_chronogram(row_str))
                if items: return Response({'items': items[:50], 'method': 'excel_generic'})

            # Parsing générique du texte extrait
            if not items:
                items = self._parse_text_to_chronogram(raw_text)
                # Remappage pour correspondre aux en-têtes par défaut
                headers = ['Heure Début', 'Heure Fin', 'Activité', 'Détails']
                mapped_items = []
                for it in items:
                    mapped = {'jour': it.get('jour', 'Jour 1')}
                    mapped['Heure Début'] = it.get('heure_debut', '')
                    mapped['Heure Fin'] = it.get('heure_fin', '')
                    mapped['Activité'] = it.get('titre', '')
                    mapped['Détails'] = it.get('description', '')
                    mapped_items.append(mapped)
                items = mapped_items

            return Response({
                'items': items[:100], 
                'headers': headers if 'headers' in locals() else ['Heure Début', 'Heure Fin', 'Activité', 'Détails'],
                'method': locals().get('method', 'generic_parsing')
            })

        except Exception as e:
            logger.exception("Erreur lors de l'extraction du chronogramme")
            return Response({'error': str(e)}, status=500)

    @action(detail=False, methods=['post'], url_path='detect-info', authentication_classes=[], permission_classes=[permissions.AllowAny])
    def detect_info(self, request):
        """Extraction intelligente réelle via Tesseract OCR."""
        image_file = request.FILES.get('image')
        if not image_file:
            return Response({'error': 'Aucune image fournie'}, status=400)
            
        logger.info(f"Début de detect_info - Fichier: {image_file.name}")
        try:
            # Ouverture et OCR
            img = Image.open(image_file)
            text = pytesseract.image_to_string(img, lang='eng')
            
            # Nettoyage et split
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            
            titre = "Événement sans titre"
            if len(lines) > 0:
                top_lines = lines[:5]
                titre = max(top_lines, key=len) if top_lines else lines[0]

            date_pattern = r'(\d{1,2}[/\\-\\.]\d{1,2}[/\\-\\.]\d{2,4})'
            dates_found = re.findall(date_pattern, text)
            
            date_debut = timezone.now() + timezone.timedelta(days=7)
            date_fin = date_debut + timezone.timedelta(hours=4)
            
            if len(dates_found) >= 1:
                try:
                    d_str = dates_found[0].replace('.', '/').replace('-', '/')
                    parts = d_str.split('/')
                    if len(parts[2]) == 2: parts[2] = "20" + parts[2]
                    date_obj = datetime.datetime.strptime("/".join(parts), "%d/%m/%Y")
                    date_debut = timezone.make_aware(date_obj, timezone.get_current_timezone())
                except Exception as de:
                    logger.error(f"Erreur parsing date début: {str(de)}")
            
            if len(dates_found) >= 2:
                try:
                    d_str = dates_found[1].replace('.', '/').replace('-', '/')
                    parts = d_str.split('/')
                    if len(parts[2]) == 2: parts[2] = "20" + parts[2]
                    date_obj = datetime.datetime.strptime("/".join(parts), "%d/%m/%Y")
                    date_fin = timezone.make_aware(date_obj, timezone.get_current_timezone())
                except Exception as de:
                    logger.error(f"Erreur parsing date fin: {str(de)}")

            lieu = "Lieu à préciser"
            lieu_keywords = ["lieu", "place", "salle", "hôtel", "stade", "église", "paroisse", "abidjan"]
            for line in lines:
                if any(kw in line.lower() for kw in lieu_keywords):
                    lieu = line.split(':')[-1].strip() if ':' in line else line
                    break

            data = {
                'titre': titre[:100],
                'date_debut': date_debut.isoformat() if hasattr(date_debut, 'isoformat') else str(date_debut),
                'date_fin': date_fin.isoformat() if hasattr(date_fin, 'isoformat') else str(date_fin),
                'lieu': lieu[:100],
                'type_evenement': 'autre',
                'description': f"Extrait de l'affiche : {text[:200]}...",
                'confidence': 0.90
            }
            if 'camp' in text.lower(): data['type_evenement'] = 'camp'
            elif 'concert' in text.lower(): data['type_evenement'] = 'concert'
            elif 'séminaire' in text.lower() or 'seminaire' in text.lower(): data['type_evenement'] = 'seminaire'
            
            return Response(data)
        except Exception as e:
            logger.exception("Erreur fatale dans detect_info OCR")
            return Response({'error': str(e)}, status=500)

    def _parse_text_to_chronogram(self, text):
        """Transforme un bloc de texte en liste d'items de chronogramme avec support multi-jours."""
        items = []
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        current_day = "Jour 1"
        
        # Regex pour les dates (07/04, 07 avril, 7 au 10...)
        date_pattern = r'(\b\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b|\b\d{1,2}[\/-]\d{1,2}[\/-]\d{2,4}\b)'
        # Keywords pour détecter un changement de jour
        day_keywords = ['jour', 'day', 'étape', 'etape', 'lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']
        
        for line in lines:
            line_lower = line.lower()
            
            # Détection de changement de jour explicite
            # Si la ligne contient un jour de la semaine ou "JOUR X" et est relativement courte
            if (any(kw in line_lower for kw in day_keywords) or re.search(date_pattern, line_lower)) and len(line) < 50:
                potential_day = line.strip(' -:|#*')
                # On évite de prendre une ligne d'activité comme un jour
                if not re.search(r'\d{1,2}[hH:]', potential_day):
                    current_day = potential_day
                    continue

            parsed = self._parse_line_for_chronogram(line)
            if parsed:
                for item in parsed:
                    item['jour'] = current_day
                    items.append(item)
        return items

    def _parse_line_for_chronogram(self, line, force_day=None):
        """Extrait une ou deux heures (début/fin) et un titre à partir d'une ligne de texte."""
        time_pattern = r'(\d{1,2}[hH:]\d{0,2})'
        matches = re.findall(time_pattern, line)
        
        if matches:
            # Heure de début
            h_debut = matches[0].lower().replace('h', ':')
            if ':' not in h_debut: h_debut += ":00"
            elif h_debut.endswith(':'): h_debut += "00"
            
            # Heure de fin (si présente)
            h_fin = ""
            if len(matches) > 1:
                h_fin = matches[1].lower().replace('h', ':')
                if ':' not in h_fin: h_fin += ":00"
                elif h_fin.endswith(':'): h_fin += "00"
            
            # Nettoyage du titre de TOUTES les heures trouvées
            title = line
            for m in matches:
                title = title.replace(m, '')
            
            title = title.strip(' -:|~—')
            if title:
                return [{
                    'heure_debut': h_debut,
                    'heure_fin': h_fin,
                    'titre': title[:100],
                    'description': '',
                    'pole': None,
                    'jour': force_day
                }]
        return []

class ChronogrammeTemplateViewSet(viewsets.ModelViewSet):
    queryset = ChronogrammeTemplate.objects.all()
    serializer_class = ChronogrammeTemplateSerializer
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    search_fields = ['nom', 'description']

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        if not user or user.is_anonymous:
            return self.queryset.none()
        # On voit ses modèles + les modèles publics/officiels (ceux des admins)
        return self.queryset.filter(Q(cree_par=user) | Q(cree_par__is_superuser=True)).distinct().order_by('-id')

    def perform_create(self, serializer):
        serializer.save(cree_par=self.request.user)

class ChronogrammeItemViewSet(viewsets.ModelViewSet):
    queryset = ChronogrammeItem.objects.all()
    serializer_class = ChronogrammeItemSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['evenement', 'pole', 'statut']
    search_fields = ['titre', 'description']
    ordering_fields = ['heure_debut']

class PoleCompetenceViewSet(viewsets.ModelViewSet):
    queryset = PoleCompetence.objects.all().order_by('nom')
    serializer_class = PoleCompetenceSerializer
    permission_classes = [permissions.AllowAny]
    pagination_class = None

class RegionViewSet(viewsets.ModelViewSet):
    queryset = Region.objects.annotate(eglise_count=Count('region_eglise')).all().order_by('nom')
    serializer_class = RegionSerializer
    permission_classes = [permissions.AllowAny]
    pagination_class = None
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    search_fields = ['nom']

class VilleViewSet(viewsets.ModelViewSet):
    queryset = Ville.objects.all().order_by('nom')
    serializer_class = VilleSerializer
    permission_classes = [permissions.AllowAny]
    pagination_class = None
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    filterset_fields = ['region']
    search_fields = ['nom']

class MouvementMaterielViewSet(viewsets.ModelViewSet):
    queryset = MouvementMateriel.objects.all().order_by('-date_mouvement')
    serializer_class = MouvementMaterielSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = {
        'type_mouvement': ['exact'],
        'evenement': ['exact'],
        'batch_id': ['exact'],
        'eglise_origine': ['exact'],
        'materiel': ['exact'],
        'date_mouvement': ['exact', 'gte', 'lte'],
    }
    search_fields = ['notes', 'batch_id', 'materiel__nom', 'eglise_origine__nom']

class FicheDefectuositeViewSet(viewsets.ModelViewSet):
    queryset = FicheDefectuosite.objects.all()
    serializer_class = FicheDefectuositeSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_fields = ['materiel', 'repare', 'niveau_gravite']

    def perform_create(self, serializer):
        serializer.save(rapporteur=self.request.user)

class ReunionDimancheViewSet(viewsets.ModelViewSet):
    queryset = ReunionDimanche.objects.all().order_by('-date_reunion')
    serializer_class = ReunionDimancheSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = pagination.PageNumberPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['eglise_hote', 'date_reunion']
    search_fields = ['ordre_du_jour', 'eglise_hote__nom']
    ordering_fields = ['date_reunion', 'heure_debut']

class RessourceFormationViewSet(viewsets.ModelViewSet):
    queryset = RessourceFormation.objects.all().order_by('-date_ajout')
    serializer_class = RessourceFormationSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = pagination.PageNumberPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['type_ressource', 'pole']
    search_fields = ['titre', 'description']
    ordering_fields = ['date_ajout', 'titre']

class DemandeFormationSGLViewSet(viewsets.ModelViewSet):
    queryset = DemandeFormationSGL.objects.all()
    serializer_class = DemandeFormationSGLSerializer
    permission_classes = [permissions.IsAuthenticated]

class ExpressionBesoinViewSet(viewsets.ModelViewSet):
    queryset = ExpressionBesoin.objects.select_related('circuit', 'eglise', 'demandeur').all().order_by('-date_demande')
    serializer_class = ExpressionBesoinSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = pagination.PageNumberPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['eglise', 'evenement', 'etape_circuit', 'demandeur']
    search_fields = ['liste_materiel', 'eglise__nom']
    ordering_fields = ['date_demande', 'estimation_budget']

    @action(detail=True, methods=['post'])
    def decider(self, request, pk=None):
        """Approuver ou refuser une demande budgétaire."""
        expr = self.get_object()
        decision = request.data.get('decision')  # 'VALIDE' | 'REFUSE' | next step
        notes = request.data.get('notes', '')
        try:
            circuit = expr.circuit
        except ValidationCircuit.DoesNotExist:
            circuit = ValidationCircuit.objects.create(besoin=expr)
        circuit.etape_actuelle = decision
        circuit.notes_decision = notes
        circuit.save()
        return Response({'status': decision, 'notes': notes})

class ValidationCircuitViewSet(viewsets.ModelViewSet):
    queryset = ValidationCircuit.objects.all()
    serializer_class = ValidationCircuitSerializer
    permission_classes = [permissions.IsAuthenticated]

# ViewSets pour la Phase 6
class FormationViewSet(viewsets.ModelViewSet):
    queryset = Formation.objects.all()
    serializer_class = FormationSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = StandardPagination

class DemandeFormationViewSet(viewsets.ModelViewSet):
    queryset = DemandeFormation.objects.select_related('eglise', 'demandeur', 'formation').all().order_by('-date_demande')
    serializer_class = DemandeFormationSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    filterset_fields = ['eglise', 'statut']
    search_fields = ['formation__titre', 'notes']

class SessionFormationViewSet(viewsets.ModelViewSet):
    queryset = SessionFormation.objects.select_related('formation').all().order_by('date_debut')
    serializer_class = SessionFormationSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['statut']


@api_view(['GET'])
def user_list(request):
    """Retourne la liste des utilisateurs (nom + id) pour les selects du frontend."""
    if not request.user.is_authenticated:
        return Response(status=403)
    User = get_user_model()
    users = User.objects.filter(is_active=True).values('id', 'first_name', 'last_name', 'email')
    return Response({'results': list(users)})

class EvenementImageViewSet(viewsets.ModelViewSet):
    queryset = EvenementImage.objects.all()
    serializer_class = EvenementImageSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]

class MaterielImageViewSet(viewsets.ModelViewSet):
    queryset = MaterielImage.objects.all()
    serializer_class = MaterielImageSerializer
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication, SessionAuthentication]

    def perform_destroy(self, instance):
        materiel = instance.materiel
        instance.delete()
        
        # Sync primary image with the first remaining gallery image
        remaining = materiel.images_materiel.all().order_by('id')
        if remaining.exists():
            materiel.image = remaining[0].image
        else:
            materiel.image = None
        materiel.save()
