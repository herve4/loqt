# export_utils.py
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import openpyxl
from docx import Document

def export_eglises_to_pdf(eglises):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Titre du document
    title = Paragraph("Liste des Églises", styles['Title'])
    
    # Préparation des données
    data = [["Nom", "Ville", "Région", "Pasteur", "Téléphone"]]
    
    for eglise in eglises:
        pasteur = eglise.pasteur.get_full_name() if eglise.pasteur else "Non assigné"
        region = eglise.region.nom if eglise.region else (eglise.ville.region.nom if eglise.ville and eglise.ville.region else "Non spécifiée")
        phone = eglise.phone if eglise.phone else "Non renseigné"
        data.append([
            eglise.nom or "Non renseigné",
            eglise.ville.nom,
            region,
            pasteur,
            phone
        ])
        print(data)
    
    # Création du tableau
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    
    # Construction du PDF
    elements = [title, table]
    doc.build(elements)
    
    buffer.seek(0)
    return buffer

def export_eglises_to_excel(eglises):
    buffer = io.BytesIO()
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Liste des Églises"
    
    # En-têtes
    sheet.append(["Nom", "Ville", "Région", "Pasteur", "Téléphone"])
    
    # Données
    for eglise in eglises:
        pasteur = eglise.pasteur.get_full_name() if eglise.pasteur else "Non assigné"
        region = eglise.region.nom if eglise.region else (eglise.ville.region.nom if eglise.ville and eglise.ville.region else "Non spécifiée")
        phone = eglise.phone if eglise.phone else "Non renseigné"
        sheet.append([
            eglise.nom or "Non renseigné",
            eglise.ville.nom,
            region,
            pasteur,
            phone
        ])
    
    workbook.save(buffer)
    buffer.seek(0)
    return buffer

def export_eglises_to_word(eglises):
    buffer = io.BytesIO()
    document = Document()
    
    # Titre du document
    document.add_heading('Liste des Églises', 0)
    
    # Création du tableau
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nom'
    hdr_cells[1].text = 'Ville'
    hdr_cells[2].text = 'Région'
    hdr_cells[3].text = 'Pasteur'
    hdr_cells[4].text = 'Téléphone'
    
    # Données
    for eglise in eglises:
        pasteur = eglise.pasteur.get_full_name() if eglise.pasteur else "Non assigné"
        region = eglise.region.nom if eglise.region else (eglise.ville.region.nom if eglise.ville and eglise.ville.region else "Non spécifiée")
        phone = eglise.phone if eglise.phone else "Non renseigné"
        row_cells = table.add_row().cells
        row_cells[0].text = eglise.nom or "Non renseigné"
        row_cells[1].text = eglise.ville.nom or "Non renseigné"
        row_cells[2].text = region
        row_cells[3].text = pasteur
        row_cells[4].text = phone
    
    document.save(buffer)
    buffer.seek(0)
    return buffer